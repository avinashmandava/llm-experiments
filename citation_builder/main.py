from langchain.prompts import PromptTemplate, ChatPromptTemplate, HumanMessagePromptTemplate
from langchain.schema import (
    AIMessage,
    HumanMessage,
    SystemMessage
)
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.output_parsers import PydanticOutputParser, OutputFixingParser
from langchain.document_loaders import UnstructuredWordDocumentLoader
from pydantic import BaseModel, Field, validator
from typing import List
from dotenv import load_dotenv
import json
import os
import re
import validators
import docx
from lxml import etree
import bibtexparser
import ast

# CONFIG
load_dotenv()
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
DATA_DIR = os.getenv('DATA_DIR')
#system_message = "You are a helpful research assistant. Your job is to take a list of footnotes from research papers and create a bibtex-style list of citations for all sources in the list of footnotes."
system_message = "You are a helpful research assistant. Your job is to take a list of footnotes from research papers and create a list of citations for all sources in the list of footnotes. The citations should follow the Chicago Manual of Style"
chat = ChatOpenAI(model_name="gpt-3.5-turbo")
batch_size = 5

class Citations(BaseModel):
    #citations: List[str] = Field(description="Python list of all bibtex-formatted citations extracted from the supplied list of footnotes")
    citations: List[str] = Field(description="Python list where each element of the list is one Chicago Manual of Style-formatted citation extracted from the supplied list of footnotes.")

def load_footnotes_from_file(file_path: str) -> List:
    doc = docx.Document(file_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return lines

def extract_footnotes(doc_path):
    footnotes = []
    doc = docx.Document(doc_path)

    for part in doc.part.related_parts.values():
        if part.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml':
            footnotes_xml = part.blob
            root = etree.fromstring(footnotes_xml)
            namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            footnote_elems = root.findall('.//w:footnote', namespace)

            for footnote_elem in footnote_elems:
                footnote_text = ''
                for elem in footnote_elem.iter():
                    if elem.tag.endswith('}t'):
                        if elem.text:
                            footnote_text += elem.text
                if footnote_text:
                    footnotes.append(footnote_text)
    return footnotes

def get_citations(footnotes: List):
  # Your code goes here

  parser = PydanticOutputParser(pydantic_object=Citations)

  prompt = PromptTemplate(
    #template = "Take the folllowing list of footnotes, extract all the source info, and construct a list of bibliographic citations. Do it the same way bibtex would do it. And present the citations to me in a python list:\n{footnotes}\n{format_instructions}.\nThere should be nothing other than the raw JSON object returned in your response. Do not add any markdown formatting or other text to your response. Do not preface the raw JSON response with any text.",
    template = "Take the folllowing list of footnotes, extract all the source info, and construct a list of bibliographic citations. Format the citations according to the Chicago Manual of Style. And present the citations to me in a python list where each element is one citation:\n{footnotes}\n{format_instructions}.\nThere should be nothing other than the raw list object returned in your response. Do not add any markdown formatting or other text to your response. Do not preface the raw list response with any text.",
    input_variables = ["footnotes"],
    partial_variables = {"format_instructions": parser.get_format_instructions()}
  )
  input = prompt.format_prompt(footnotes=footnotes)
  messages = [
    SystemMessage(content=system_message),
    HumanMessage(content=input.to_string())
  ]
  # Sometimes the response comes back with markdown formatting or is prefaced with "json", so we strip it out
  output = chat(messages).content.strip("```").strip("json").strip().strip("```").strip()
  # We deal with escape characters with a litreal eval that converts the string to a python dict object
  print(output)
  output = ast.literal_eval(output)
  return output["citations"]

def main():
  file_name="Footnotes.docx"
  file_path = f"{DATA_DIR}/{file_name}"
  # Get the footnotes
  #footnotes = extract_footnotes(file_path)
  footnotes = load_footnotes_from_file(file_path)

  print(f"Found {len(footnotes)} footnotes in {file_name}")
  # One batch at a time
  output_filename = f"{DATA_DIR}/{file_name}_citations.txt"
  with open(output_filename, 'w+') as f:
    f.write('')
  citations = []
  start = 0
  end = start + batch_size
  with open(output_filename, "a") as f:
    for i in range(1040, len(footnotes), batch_size):
      print(f"Processing from {i} to {i+batch_size} of {len(footnotes)} footnotes")
      # Process the next batch of 5 data points, or any remaining points if there are fewer than 5
      try:
        results = get_citations(footnotes[i:i+batch_size])
      except:
        print(f"Error processing batch starting at {i} of {len(footnotes)} footnotes. Trying again...")
        results = get_citations(footnotes[i:i+batch_size])

      # Write the results for the current batch to the output file
      for result in results:
        f.write(result + '\n')

    # If there are any remaining data points that didn't form a complete batch of 5, process and write them too
    if len(footnotes) % batch_size != 0:
      results = get_citations(footnotes[-(len(footnotes) % batch_size):])
      for result in results:
        f.write(result + '\n')

if __name__ == "__main__":
    main()
